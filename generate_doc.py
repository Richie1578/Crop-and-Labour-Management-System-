from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('FarmTrack: Advanced Crop & Labour Management', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('A comprehensive digital ecosystem for the modern Zambian farmer.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break
    
    # 1. Step-by-Step Functionality
    doc.add_heading('1. What the App Does: A Step-by-Step Guide', level=1)
    
    steps = [
        ("Phase 1: Digital Onboarding & GPS Mapping", 
         "The journey begins with the farmer setting up their profile. Using the 'GPS Landscape' tool, the farmer can visualize their land via high-resolution satellite imagery. By clicking the boundaries of their fields, the app calculates the exact acreage. This is critical for Zambia, where miscalculating field size often leads to wasting expensive fertilizer."),
        ("Phase 2: Inventory & Labour Registration", 
         "The farmer catalogs all workers, their roles (Supervisor, Tractor Operator, etc.), and their daily wage rates. Simultaneously, all farm inputs—seeds, fertilizers like Urea and D-Compound, and chemicals—are entered into the digital inventory with their purchase prices."),
        ("Phase 3: Real-Time Operational Logging", 
         "As the season progresses, tasks are assigned (e.g., 'Apply Top Dress to North Field'). Workers' attendance is logged daily. The app automatically calculates payroll totals, ensuring transparency and reducing disputes."),
        ("Phase 4: Intelligent Monitoring & AI Advice", 
         "The farmer uses the AI Crop Advisor to monitor their crops. If a pest appears, the farmer takes a photo. The AI (powered by Claude) analyzes the image, identifies the pest (like Fall Armyworm), and suggests the exact treatment and application rates localized for Zambia."),
        ("Phase 5: Financial Realization & Analytics", 
         "At harvest, the farmer logs the yield and sale price. The system automatically subtracts all labour and input costs to show a 'Net Profit' report. This allows the farmer to see exactly which crop made money and which field was a loss.")
    ]
    
    for title, desc in steps:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(title)
        run.bold = True
        doc.add_paragraph(desc)

    # 2. How it Works
    doc.add_heading('2. How it Works: Technical & Operational Overview', level=1)
    doc.add_paragraph(
        "FarmTrack is a high-performance Single Page Application (SPA) designed for the unique constraints of rural farming environments. "
        "Technically, it utilizes a 'Clean Stack' (HTML5, Vanilla CSS, and JavaScript) to ensure lightning-fast load times even on older smartphones or weak 3G/4G connections."
    )
    
    tech_points = [
        ("Geospatial Layer", "Integrated with Leaflet.js and Esri World Imagery for precise satellite mapping."),
        ("Intelligence Layer", "Connected to Claude AI via Supabase Edge Functions for advanced image recognition and agricultural logic."),
        ("Data Resilience", "Uses an 'Offline-First' architecture. Data is saved locally in the browser and synced to the cloud whenever a signal is available."),
        ("Localization Engine", "Features a built-in Zambian Planting Calendar and supports multiple local languages including Nyanja and Bemba.")
    ]
    
    for k, v in tech_points:
        p = doc.add_paragraph('', style='List Bullet')
        run = p.add_run(f"{k}: ")
        run.bold = True
        p.add_run(v)

    # 3. Solving Zambian Problems
    doc.add_heading('3. Problems Solved in the Zambian Context', level=1)
    
    problems = [
        ("Eliminating the 'Input Waste' Gap", 
         "In Zambia, many farmers purchase inputs based on estimates. Miscalculating a 2-hectare field as 2.5 hectares results in a 25% waste of fertilizer—money that could have been profit. FarmTrack’s GPS mapper provides the precision needed to optimize every bag of fertilizer."),
        ("Professionalizing Farm Records", 
         "Traditionally, records are kept in paper notebooks which are easily lost or damaged. FarmTrack provides a permanent, searchable digital history of every season, worker, and cost."),
        ("Bridging the Agronomy Gap", 
         "Professional extension officers are often spread thin. The AI Advisor provides 24/7 expert advice on crop diseases and planting windows, tailored specifically to the Zambian climate and soil types.")
    ]
    
    for title, desc in problems:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.underline = True
        doc.add_paragraph(desc)

    # 4. Competitive Advantage
    doc.add_heading('4. Why FarmTrack is Superior to Existing Systems', level=1)
    doc.add_paragraph(
        "Unlike generic global agricultural apps, FarmTrack is 'Zambia-First'. It doesn't just track crops; it understands the Zambian economy. "
        "It uses Zambian Kwacha (ZMW) by default, integrates with the FISP (Farmer Input Support Programme) planting windows, and understands local challenges like armyworm infestations."
    )
    doc.add_paragraph(
        "While other apps focus only on large-scale logistics, FarmTrack merges Labour, Finance, and Agronomy into a single, simple interface that works in a farmer's hand, not just in an office."
    )

    # 5. The Personal Perspective: Benefits for the Farmer
    doc.add_heading('5. The Benefits: Why This Makes Life Easier', level=1)
    
    benefits = [
        "Data-Driven Decisions: Know exactly where your money is going. If Soya is more profitable than Maize this season, you will have the data to prove it.",
        "Financial Inclusion: The professional reports generated (CSV/PDF) can be presented to banks or micro-lenders as a 'Business Case' for seasonal loans.",
        "Administrative Freedom: Stop spending hours every Friday night calculating worker wages. The system does it for you instantly.",
        "Increased Yields: Through precision measurement and AI-guided disease prevention, farmers can see significant increases in harvest quality and quantity."
    ]
    
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    doc.save('FarmTrack_System_Overview.docx')
    print("Document 'FarmTrack_System_Overview.docx' created successfully.")

if __name__ == "__main__":
    create_doc()
